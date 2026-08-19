from __future__ import annotations

import os
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "2026-08-19_thailand-line-oa-plan"
OUT.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT / "泰国_LINE-OA课前运营与销售分配实验方案_v1.0.docx"
MD_PATH = OUT / "泰国_LINE-OA课前运营与销售分配实验方案_v1.0.md"
FLOW_MAIN = OUT / "flow_main.png"
FLOW_ASSIGN = OUT / "flow_assignment.png"

SKILL_ROOT = Path(
    "/Users/liuniuniu/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.818.11542/skills/documents"
)

BLUE = "2E74B5"
NAVY = "163A5F"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E8F5EE"
PALE_YELLOW = "FFF5D9"
PALE_RED = "FDECEC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1F2937"
WHITE = "FFFFFF"
GREEN = "138A5B"
ORANGE = "B36B00"
RED = "B42318"
DOC_FONT = "Heiti SC"


def rgb(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str)


def set_run_font(run, *, name=DOC_FONT, east_asia=DOC_FONT, size=11,
                 color=DARK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, *, bold=False, color=DARK, size=9.2,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY,
                      size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(str(value)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), size=font_size, align=align)
        set_row_cant_split(table.rows[-1])
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_paragraph(doc, text="", *, bold=False, color=DARK, size=11,
                  italic=False, after=6, before=0, align=None,
                  keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_callout(doc, label, text, *, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=DARK)
    set_table_geometry(table, [9360], indent_dxa=120)
    add_paragraph(doc, "", after=2)
    return table


def add_picture_with_caption(doc, path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MID_GRAY, italic=True)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def round_rect(draw, xy, fill, outline, radius=24, width=3):
    if isinstance(fill, str) and not fill.startswith("#"):
        fill = "#" + fill
    if isinstance(outline, str) and not outline.startswith("#"):
        outline = "#" + outline
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def load_font(size, bold=False):
    path = "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc"
    return ImageFont.truetype(path, size=size)


def wrapped_lines(text, font, max_width, draw):
    lines = []
    for para in text.split("\n"):
        current = ""
        for char in para:
            candidate = current + char
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = char
        if current:
            lines.append(current)
    return lines


def draw_box(draw, xy, text, fill, outline=BLUE, font_size=34, text_color=DARK, bold=False):
    round_rect(draw, xy, fill, outline, radius=22, width=3)
    if isinstance(text_color, str) and not text_color.startswith("#"):
        text_color = "#" + text_color
    font = load_font(font_size, bold)
    x1, y1, x2, y2 = xy
    lines = wrapped_lines(text, font, x2 - x1 - 34, draw)
    line_h = font_size + 13
    total_h = len(lines) * line_h - 8
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        tx = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((tx, y), line, font=font, fill=text_color)
        y += line_h


def arrow(draw, start, end, color="#667085", width=5):
    draw.line([start, end], fill=color, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    for delta in (2.55, -2.55):
        p = (end[0] + size * math.cos(angle + delta), end[1] + size * math.sin(angle + delta))
        draw.line([end, p], fill=color, width=width)


def label(draw, pos, text, color=MID_GRAY, size=27, bold=True):
    draw.text(pos, text, font=load_font(size, bold), fill="#" + color)


def build_main_flow():
    img = Image.new("RGB", (2800, 1900), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(54, True)
    d.text((100, 60), "泰国 LINE-OA 课前运营主流程", font=title_font, fill="#" + NAVY)
    d.text((100, 130), "OA 覆盖全渠道；外呼仅对配置渠道启用，暂定泰国时间 08:00–21:00 按 +5/+15/+30 分钟重试", font=load_font(28), fill="#" + MID_GRAY)

    nodes = {
        "reg": (100, 240, 720, 380),
        "cta": (900, 240, 1720, 380),
        "check": (1900, 240, 2700, 380),
        "call": (1900, 540, 2700, 720),
        "not_follow": (1900, 840, 2700, 1000),
        "oa": (100, 540, 900, 700),
        "service": (100, 820, 900, 1000),
        "book": (100, 1140, 720, 1280),
        "remind": (900, 1140, 1720, 1280),
        "attend": (1900, 1140, 2700, 1280),
        "first_end": (1900, 1430, 2700, 1590),
        "resched": (1000, 1430, 1720, 1590),
        "fail_end": (100, 1690, 820, 1830),
        "second": (1000, 1690, 1800, 1830),
        "end": (1980, 1690, 2700, 1830),
    }
    draw_box(d, nodes["reg"], "全渠道用户注册完成", "#" + PALE_BLUE, bold=True)
    draw_box(d, nodes["cta"], "完成页一键关注 OA\n移动端深链接；桌面端二维码", "#" + PALE_GREEN)
    draw_box(d, nodes["check"], "注册 +5 分钟\n查询该用户 OA 关注状态", "#" + PALE_YELLOW, outline=ORANGE, bold=True)
    draw_box(d, nodes["call"], "配置渠道且在时段内：AI 电话\n未接通则注册后 +15 / +30 分钟重试", "#" + PALE_RED, outline=RED)
    draw_box(d, nodes["not_follow"], "三次未接通、已超时段或仍未关注\n记录原因并结束本轮补促", "#" + LIGHT_GRAY, outline=MID_GRAY)
    draw_box(d, nodes["oa"], "已关注 OA\n绑定 user_id / 手机号 / LINE userId", "#" + PALE_GREEN, outline=GREEN, bold=True)
    draw_box(d, nodes["service"], "OA 自动化：问好 → 轻量挖需 → 简单问答\n未预约则协助预约；过程回写 CRM", "#" + PALE_BLUE)
    draw_box(d, nodes["book"], "完成首次体验课预约", "#" + PALE_GREEN, outline=GREEN)
    draw_box(d, nodes["remind"], "首次课前提醒 + 入室引导", "#" + PALE_BLUE)
    draw_box(d, nodes["attend"], "首次体验课是否出席？", "#" + PALE_YELLOW, outline=ORANGE, bold=True)
    draw_box(d, nodes["first_end"], "出席：OA SOP 结束\n按渠道规则完成销售归属/承接", "#" + PALE_GREEN, outline=GREEN)
    draw_box(d, nodes["resched"], "缺席：发起 1 次二次预约", "#" + PALE_RED, outline=RED)
    draw_box(d, nodes["fail_end"], "二约失败：本轮结束\n记录原因与最终状态", "#" + LIGHT_GRAY, outline=MID_GRAY)
    draw_box(d, nodes["second"], "二约成功：第二次课前提醒\n并确认出席", "#" + PALE_BLUE)
    draw_box(d, nodes["end"], "第二次出席或再次缺席\nLINE-OA SOP 结束", "#" + PALE_GREEN, outline=GREEN, bold=True)

    arrow(d, (720, 310), (900, 310))
    arrow(d, (1720, 310), (1900, 310))
    arrow(d, (2300, 380), (2300, 540))
    label(d, (2325, 430), "未关注")
    arrow(d, (1900, 630), (900, 630))
    label(d, (1230, 580), "后续关注成功")
    arrow(d, (2300, 720), (2300, 840))
    label(d, (2325, 755), "停止条件")
    arrow(d, (1900, 310), (500, 540))
    label(d, (1180, 400), "已关注")
    arrow(d, (500, 700), (500, 820))
    arrow(d, (500, 1000), (410, 1140))
    arrow(d, (720, 1210), (900, 1210))
    arrow(d, (1720, 1210), (1900, 1210))
    arrow(d, (2300, 1280), (2300, 1430))
    label(d, (2325, 1330), "出席")
    arrow(d, (1900, 1210), (1720, 1510))
    label(d, (1770, 1340), "缺席")
    arrow(d, (1000, 1510), (460, 1690))
    label(d, (660, 1575), "失败")
    arrow(d, (1360, 1590), (1400, 1690))
    label(d, (1420, 1620), "成功")
    arrow(d, (1800, 1760), (1980, 1760))
    img.save(FLOW_MAIN, quality=95)


def build_assignment_flow():
    img = Image.new("RGB", (2800, 2500), "white")
    d = ImageDraw.Draw(img)
    d.text((100, 60), "销售分配机制：当前实验与长期目标", font=load_font(54, True), fill="#" + NAVY)
    d.text((100, 130), "非安卓用户按稳定随机规则进入 30% 直分组 / 70% 抢新组；安卓维持出席后分配", font=load_font(28), fill="#" + MID_GRAY)
    draw_box(d,(950,230,1850,380),"泰国用户注册", "#"+PALE_BLUE, bold=True)
    draw_box(d,(950,500,1850,650),"渠道类型？", "#"+PALE_YELLOW, outline=ORANGE, bold=True)

    draw_box(d,(120,800,860,980),"安卓渠道\n出席前不分配销售", "#"+PALE_RED, outline=RED, bold=True)
    draw_box(d,(1030,800,1770,980),"非安卓 30%\n注册后立即直分销售", "#"+PALE_GREEN, outline=GREEN, bold=True)
    draw_box(d,(1940,800,2680,980),"非安卓 70%\n进入抢新池，暂不归属", "#"+PALE_BLUE, bold=True)

    draw_box(d,(120,1120,860,1320),"LINE-OA 负责课前运营\n首次出席后系统分配销售", "#"+PALE_BLUE)
    draw_box(d,(1030,1120,1770,1360),"LINE-OA 继续自动运营\n销售用手机号搜索用户并以私人 LINE 单独沟通\n两条会话通过 CRM 同步", "#"+PALE_GREEN, outline=GREEN)
    draw_box(d,(1940,1120,2680,1360),"LINE-OA 继续自动运营\n真人销售电话打通后才归属\n归属后可用私人 LINE 沟通", "#"+PALE_BLUE)

    draw_box(d,(580,1570,2220,1770),"当前阶段共同限制：OA 与私人 LINE 是两条独立会话\n需统一触达规则、用户口径、消息记录和冲突处理，避免重复催促", "#"+PALE_YELLOW, outline=ORANGE, bold=True)
    draw_box(d,(580,1960,2220,2190),"长期目标：非安卓逐步扩大直分比例至 100%\n每位销售配置 LINE-OA 坐席，在同一 OA 会话中与机器人协作\n机器人做标准化服务；真人负责复杂问题、报价与成交", "#"+PALE_GREEN, outline=GREEN, bold=True)

    arrow(d,(1400,380),(1400,500))
    arrow(d,(1120,650),(490,800)); label(d,(600,690),"安卓")
    arrow(d,(1400,650),(1400,800)); label(d,(1420,700),"非安卓随机 30%")
    arrow(d,(1680,650),(2310,800)); label(d,(1970,690),"非安卓随机 70%")
    arrow(d,(490,980),(490,1120))
    arrow(d,(1400,980),(1400,1120))
    arrow(d,(2310,980),(2310,1120))
    arrow(d,(490,1320),(900,1570))
    arrow(d,(1400,1360),(1400,1570))
    arrow(d,(2310,1360),(1900,1570))
    arrow(d,(1400,1770),(1400,1960))
    img.save(FLOW_ASSIGN, quality=95)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Title", 23, NAVY, 0, 4),
        ("Subtitle", 13, MID_GRAY, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        st = styles[name]
        st.font.name = DOC_FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:cs"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = DOC_FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:cs"), DOC_FONT)
        st._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("泰国 LINE-OA 课前运营与销售分配实验方案")
    set_run_font(r, size=23, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("覆盖全渠道注册、OA 关注补促、课前自动化、缺席二约与销售分配机制演进")
    set_run_font(r, size=13, color=MID_GRAY)
    for label_text, value in (
        ("市场：", "泰国"),
        ("版本：", "v1.0"),
        ("日期：", "2026-08-19"),
        ("状态：", "跨部门评审稿"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label_text)
        set_run_font(r, size=10.5, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("AI-CC｜Thailand LINE-OA Project")
    set_run_font(r, size=8.5, color=MID_GRAY, bold=True)
    add_page_number(section.footer.paragraphs[0])

    props = doc.core_properties
    props.title = "泰国 LINE-OA 课前运营与销售分配实验方案"
    props.subject = "泰国 LINE-OA 全渠道课前运营与销售分配 A/B 测试"
    props.author = "AI-CC Project Team"
    props.keywords = "Thailand, LINE Official Account, AI-CC, sales assignment, A/B test"

    add_title_block(doc)
    add_callout(doc, "核心结论", "先把 LINE-OA 建成所有渠道注册后的统一课前运营入口，再通过非安卓渠道 30% 直分 / 70% 抢新的随机实验验证销售归属机制。当前阶段允许 OA 与销售私人 LINE 分线沟通，但必须由 CRM 统一记录与控频；长期升级到销售 OA 坐席后，再扩大为 100% 注册即直分。", fill=PALE_GREEN, accent=GREEN)

    add_heading(doc, "1. 项目目标与范围", 1)
    add_paragraph(doc, "本方案面向泰国市场所有注册渠道，覆盖用户注册完成至体验课流程结束之间的建联、问好、轻量挖需、预约、提醒、缺席二约和销售归属。LINE-OA 的服务边界止于首次出席，或一次二约后的第二次出席/再次缺席。课后报价、成交与长期经营由真人销售负责。")
    add_table(doc,
              ["目标层级", "本阶段要解决的问题", "成功信号"],
              [
                  ("第一目标", "提升注册用户 LINE-OA 关注率，建立可持续触达入口", "关注率及各触点增量提升"),
                  ("第二目标", "用 OA 自动化提升预约、出席与缺席二约效率", "预约率、出席率、二约成功率提升"),
                  ("第三目标", "验证非安卓用户直接分配销售是否优于抢新", "注册至付费、出席、人效改善且无体验风险"),
                  ("长期目标", "100% 注册即直分，每位销售拥有 OA 坐席并与机器人协作", "同一会话协作、统一数据、可规模化运营"),
              ], [1500, 4740, 3120])

    add_heading(doc, "2. 已确认的业务规则", 1)
    for item in (
        "市场范围：泰国，覆盖安卓与非安卓等所有注册渠道。",
        "身份能力：第三方工具支持注册 user_id / 手机号与 LINE userId 绑定，并可查询具体用户的 OA 关注状态。",
        "关注补促：注册完成页先引导关注；注册 5 分钟后查询状态。外呼仅对运营配置的较高转化渠道启用；外呼窗口暂定为泰国时间 08:00–21:00，注册时间和调用时点均须在窗口内，按注册后 +5 / +15 / +30 分钟最多尝试 3 次，接通或已关注即停止，超出窗口的尝试直接跳过、不顺延。",
        "安卓渠道：出席前不分配真人销售；首次或第二次体验课出席后再分配。",
        "非安卓 30% 组：注册后立即分配销售；OA 继续自动运营；销售通过手机号主动搜索用户，并使用私人 LINE 单独沟通。",
        "非安卓 70% 组：维持抢新；真人销售实际打通电话后才获得用户归属。",
        "缺席规则：首次缺席后只发起 1 次二次预约；二约成功后继续第二次课前提醒；第二次出席或再次缺席后结束 OA SOP。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "重要边界", "LINE-OA 的自动化范围是“注册后至体验课流程结束”，不是长期营销机器人。体验课出席后不再继续课前消息；二次缺席后也不进行第三次预约。", fill=PALE_YELLOW, accent=ORANGE)

    add_heading(doc, "3. 端到端业务流程", 1)
    add_picture_with_caption(doc, FLOW_MAIN, "图 1｜泰国 LINE-OA 从注册、关注补促到二约结束的主流程", width=6.25)
    add_heading(doc, "3.1 LINE-OA 自动化服务内容", 2)
    add_table(doc,
              ["阶段", "LINE-OA 动作", "结束/转移条件"],
              [
                  ("关注后欢迎", "即时问好，说明身份与可提供的帮助；读取注册与预约状态", "进入挖需或预约确认"),
                  ("轻量挖需", "学习对象、年龄/年级、英语水平/经历、学习目的、可上课时间", "字段达到最小可用集，不追求一次填满"),
                  ("预约协助", "未预约用户获取可用时段并完成预约；已预约用户确认时间", "预约成功"),
                  ("简单问答", "处理上课入口、时间、设备、老师/课程基础说明；复杂价格与成交问题转真人规则", "问题解决或记录待跟进"),
                  ("课前提醒", "首次及二约后的第二次体验课均执行提醒和入室引导", "到课时间"),
                  ("缺席二约", "首次缺席后发起一次二次预约；成功后重新进入提醒", "二约失败、第二次出席或再次缺席"),
              ], [1350, 5520, 2490])

    add_heading(doc, "4. 提升 LINE-OA 关注率的建议", 1)
    add_paragraph(doc, "LINE 官方能力可以解决“如何让用户更容易关注”和“如何判断是否已关注”；AI 电话与短信则是业务侧补促手段。建议采用分层触达，而不是依赖电话播报 OA ID。")
    add_number(doc, "注册完成页作为主入口：移动端放置高辨识度的一键关注按钮，直接打开 OA 添加好友/资料页；桌面端同时展示二维码。页面保留明确价值，如预约协助、课前提醒和上课入口。")
    add_number(doc, "优先评估 LINE Login add friend option：若注册或登录链路可以使用 LINE Login，可在授权时展示添加 OA 选项，并用官方 Friendship Status API 返回关注状态。")
    add_number(doc, "注册 +5 分钟查询：按绑定后的 user_id / LINE userId 查询；已关注者立即进入欢迎 SOP，未关注者仅在渠道白名单和外呼时段均符合时进入 AI 电话任务。")
    add_number(doc, "AI 电话最多三次：外呼窗口暂定为泰国时间 08:00–21:00。窗口内注册的配置渠道用户，在注册后 +5 分钟首呼；未接通则 +15、+30 分钟重试。每次呼叫前重新查询关注状态，接通或已关注即停止；调用时点超过 21:00 则跳过，不顺延到次日。")
    add_number(doc, "AI 电话只做一件事：确认用户是否方便接收 LINE 学习服务，并提示即将发送可点击链接。电话中不要求用户逐字符记录 OA ID。")
    add_number(doc, "短信作为操作载体：通话中或通话后立即发送带 OA 深链接的短信；链接先经过自有可追踪跳转页再打开 LINE，以记录点击与触点归因。短信发送需遵守当地通信同意、退订与频控要求。")
    add_number(doc, "OA ID 仅作兜底：当深链接或二维码无法使用时，才让用户在 LINE 中搜索官方 Basic ID/Premium ID。")
    add_callout(doc, "为什么不建议只报 OA 号码", "语音场景下容易听错、记错和中途退出；可点击深链接把操作从多步搜索缩短为打开 LINE 后确认关注。销售通过手机号搜索用户同样不保证成功，受用户隐私设置和账号绑定状态影响，必须单独统计搜索成功率。", fill=PALE_RED, accent=RED)

    add_heading(doc, "4.1 关注率漏斗与归因", 2)
    add_table(doc,
              ["节点", "建议事件", "用途"],
              [
                  ("完成页曝光", "oa_cta_impression", "确认用户看到关注入口"),
                  ("点击添加", "oa_cta_click", "衡量页面引导效率"),
                  ("关注成功", "oa_follow_success", "核心结果；记录关注时间与来源"),
                  ("5 分钟检测", "oa_follow_check_5m", "区分自然关注与补促对象"),
                  ("AI 电话", "oa_follow_call_attempt / sequence / result", "区分 +5/+15/+30 分钟轮次、跳过原因、接通与停止原因"),
                  ("短信点击", "oa_sms_link_click", "衡量通话后操作完成"),
                  ("取消关注", "oa_unfollow", "体验与消息质量护栏"),
              ], [1800, 3300, 4260])

    add_heading(doc, "5. 销售分配机制与 A/B 测试", 1)
    add_picture_with_caption(doc, FLOW_ASSIGN, "图 2｜当前渠道分配规则、非安卓 A/B 测试与长期坐席目标")
    add_heading(doc, "5.1 实验组定义", 2)
    add_table(doc,
              ["人群", "分配时点", "OA 与真人销售关系", "目的"],
              [
                  ("安卓", "首次或二次体验课出席后", "出席前由 OA 独立运营；出席后分配真人", "控制低质量渠道的人力投入"),
                  ("非安卓 A：30%", "注册后立即分配", "OA 自动运营；销售用手机号搜索并通过私人 LINE 分线沟通", "验证前置归属能否提升响应、出席和付费"),
                  ("非安卓 B：70%", "真人销售电话打通后", "打通前 OA 独立运营；打通归属后销售可用私人 LINE 沟通", "保留当前抢新机制作为对照"),
              ], [1350, 1890, 3990, 2130])
    add_callout(doc, "实验原则", "30%/70% 必须在“符合条件的非安卓新注册用户”中由系统稳定随机分组，并固定用户分组，不允许销售手动挑选。按渠道、注册日期/时段等做分层或协变量校正，避免流量质量差异把结果带偏。", fill=PALE_YELLOW, accent=ORANGE)

    add_heading(doc, "5.2 A/B 测试指标", 2)
    add_paragraph(doc, "OA 关注率是全项目的核心前置指标；销售分配实验则不能只看关注率，因为 30%/70% 的差异发生在销售归属机制。建议使用以下指标层级：")
    add_table(doc,
              ["层级", "指标", "判断用途"],
              [
                  ("项目主指标", "注册后 5 分钟 / 30 分钟 / 24 小时 OA 关注率", "判断统一建联入口是否建立"),
                  ("分配实验主指标", "注册至付费转化率、每注册用户收入（观察窗固定）", "判断直分是否创造最终经营价值"),
                  ("领先指标", "真人电话接通率、私人 LINE 建联率、预约率、首次/最终出席率", "定位机制影响路径"),
                  ("效率指标", "每销售有效联系数、人均归属数、首触达时长、每单人工触达量", "判断人力与分配效率"),
                  ("护栏指标", "重复触达投诉率、OA 屏蔽率、退订率、未联系积压、销售负载离散度", "避免以转化换取用户体验或组织失衡"),
              ], [1500, 4980, 2880])
    add_paragraph(doc, "统计建议：上线前锁定观察窗、最小样本量与停止规则；优先报告绝对差、相对提升和置信区间。若样本量不足，不以短期波动直接扩到 100% 直分。", italic=True, color=MID_GRAY, size=10)

    add_heading(doc, "6. 核心指标口径", 1)
    add_table(doc,
              ["指标", "定义", "关键拆分"],
              [
                  ("OA 关注率", "在统计窗口内成功绑定且 friendFlag=true 的去重注册用户数 ÷ 符合条件的去重注册用户数", "渠道、设备、完成页版本、自然/电话/短信触点"),
                  ("5 分钟自然关注率", "AI 电话触发前已关注人数 ÷ 符合条件注册人数", "衡量完成页本身"),
                  ("补促增量关注率", "5 分钟未关注、经电话/短信后在窗口内关注人数 ÷ 5 分钟未关注人数", "接通/未接通、短信点击/未点击"),
                  ("OA 有效建联率", "关注后至少完成一次有效互动或进入自动化状态人数 ÷ 已关注人数", "关注但沉默、欢迎消息送达"),
                  ("首次预约率", "完成首次预约人数 ÷ 注册人数", "自主预约 / OA 协助预约"),
                  ("最终出席率", "首次或二次体验课至少出席一次人数 ÷ 注册人数", "首次出席、二约后出席"),
                  ("二约成功率", "首次缺席后完成二次预约人数 ÷ 首次缺席且进入二约人数", "成功时长、时段"),
              ], [1800, 4980, 2580], font_size=8.9)
    add_callout(doc, "口径注意", "分母要排除测试账号、重复注册、无效手机号和技术故障用户；同一用户跨渠道重复注册需按统一 user_id 去重。关注成功以第三方工具返回的有效关注状态为准，不以“点击按钮”代替。", fill=PALE_BLUE, accent=BLUE)

    add_heading(doc, "7. 跨部门职责与意义", 1)
    add_paragraph(doc, "本项目不是单一工具接入，而是市场入口、产研系统、运营策略、销售机制与数据评估共同组成的经营流程。每个部门需要对可验收交付物负责。")
    add_table(doc,
              ["部门/角色", "需要完成的事情", "业务意义", "主要交付/指标"],
              [
                  ("项目负责人 / PMO", "统一范围、里程碑、决策记录、跨部门依赖和灰度节奏；组织周度复盘", "防止 OA、外呼和分配机制各自优化但端到端失效", "项目看板、风险清单、Go/No-Go 决策"),
                  ("市场", "按渠道传递 source/campaign；改造注册完成页价值表达与关注入口；保证流量分组不被活动干扰", "提高自然关注率，并让结果可归因到真实渠道质量", "完成页曝光/点击/关注率，渠道有效注册"),
                  ("产品/产研", "接入第三方 OA、身份绑定与状态查询；实现渠道白名单、+5/+15/+30 分钟外呼、泰国时区窗口、OA SOP、预约/提醒/二约、分配随机化和退出开关", "把流程变成稳定、可追踪、可停止的系统能力", "事件成功率、任务准时率、异常率、分组一致性"),
                  ("前端运营 / 用户运营", "设计泰语欢迎、挖需、FAQ、预约、提醒、二约话术；配置频控、安静时段、异常兜底；抽检会话", "保证自动化不是机械催促，而是真正帮助用户完成上课", "回复率、预约率、出席率、屏蔽/投诉率"),
                  ("销售运营", "定义安卓、30%、70% 的归属规则和销售容量；制定私人 LINE 与 OA 的协同 SOP；处理冲突归属", "确保实验可执行，避免重复触达、抢单争议和负载不均", "归属时长、负载离散度、冲突率、SOP 执行率"),
                  ("真人销售", "按组别执行电话与私人 LINE 触达；查看 CRM 中 OA 摘要；记录打通、搜索/添加和沟通结果；出席后承接转化", "把 AI 沉淀转化为高质量人工跟进，而不是重复问询", "接通率、私人 LINE 建联率、出席/付费、人均产能"),
                  ("数据/BI", "建设统一漏斗、实验分组校验和渠道看板；锁定指标口径、观察窗、显著性和异常监控", "区分自然增长、补促增量和销售机制的真实因果效果", "数据完整率、实验平衡性、周期分析报告"),
                  ("法务/合规/信息安全", "审核 AI 电话与短信同意、隐私告知、数据跨系统传输、保存周期、权限与退订/屏蔽机制", "降低个人信息、通信骚扰与第三方平台使用风险", "评审结论、权限矩阵、审计记录、事件响应"),
                  ("供应商/第三方工具", "保障绑定、状态查询、消息、Webhook、坐席和日志能力；提供 SLA、错误码和对账机制", "避免关键链路成为不可观测黑盒", "接口 SLA、Webhook 成功率、故障恢复时长"),
              ], [1350, 3300, 2580, 2130], font_size=8.4)

    add_heading(doc, "8. 系统与数据需求", 1)
    add_heading(doc, "8.1 必要状态字段", 2)
    add_table(doc,
              ["数据域", "必要字段/状态", "用途"],
              [
                  ("身份", "user_id、phone_hash/手机号权限字段、LINE userId、绑定时间、绑定来源", "唯一识别与跨系统对账"),
                  ("关注", "follow_status、首次关注时间、最近检测时间、unfollow 时间、关注触点", "关注率与补促归因"),
                  ("渠道/实验", "source、campaign、device/channel、experiment_group、assignment_mode、randomization_version", "实验稳定性与渠道分析"),
                  ("销售归属", "sales_id、assigned_at、assignment_reason、human_call_connected_at、private_line_search/add_result", "区分直分与打通后归属"),
                  ("预约/出席", "booking_no、lesson_time、reminder_status、attendance_status、reschedule_count、final_end_reason", "驱动 SOP 和终止条件"),
                  ("沟通", "AI call result、SMS send/click、OA conversation summary、last_contact_at、next_action、exception_code", "统一视图、控频与人工承接"),
              ], [1500, 5160, 2700], font_size=8.8)

    add_heading(doc, "8.2 必要控制机制", 2)
    for item in (
        "幂等：同一注册事件只能创建一组 +5 / +15 / +30 分钟检测与外呼任务；Webhook 重试不得重复发送欢迎或提醒。",
        "外呼窗口：使用泰国时区判断，当前暂定 08:00–21:00；仅运营配置渠道启用，注册时间和实际调用时间均须在窗口内，超窗任务直接跳过且不顺延。",
        "重试停止：每次呼叫前重新查询关注状态；用户已关注、任一轮已接通或已达到第三轮时，取消剩余外呼任务。",
        "状态机：关注、预约、出席、二约、结束必须有明确状态迁移，任一结束状态都要停止后续课前消息。",
        "控频：OA、AI 电话、短信和真人销售共享 last_contact_at 与 next_action；避免同一小时内多渠道连续催促。",
        "人工可见：销售在 CRM 中看到 OA 摘要、已问字段、预约状态和最近动作，不重复询问用户。",
        "实验冻结：用户首次进入实验后组别固定；跨端、重复登录和重复注册需按统一用户键处理。",
        "降级：第三方状态查询或消息能力异常时暂停自动外呼/发送，进入重试队列并告警，不把技术失败计为用户未关注。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "9. 当前阶段的人机协同规则", 1)
    add_paragraph(doc, "由于尚未购买销售 LINE-OA 坐席，30% 直分组和 70% 打通后归属用户可能同时存在 OA 会话与销售私人 LINE 会话。两条会话无法天然共享上下文，因此需要明确消息边界。")
    add_table(doc,
              ["场景", "LINE-OA 负责", "真人销售负责", "冲突处理"],
              [
                  ("常规课前服务", "挖需、FAQ、预约、提醒、缺席二约", "查看摘要，原则上不重复发送同类提醒", "CRM next_action 为唯一动作依据"),
                  ("销售已建立私人 LINE", "继续执行标准 SOP，消息中保持官方服务身份", "建立关系、回答非标准问题、推动出席或付费意愿", "销售记录最近沟通，OA 在冲突窗口内抑制营销性消息"),
                  ("价格/报名/优惠", "给出轻量说明并记录需求，不做未授权报价", "负责报价、方案、异议与成交", "明确转人工标签和 SLA"),
                  ("用户要求停止", "立即停止 OA 主动消息并记录", "停止私人 LINE 主动触达或按合规规则处理", "跨系统同步 opt-out"),
                  ("用户已出席", "立即结束课前 SOP", "接管后续转化与长期服务", "出席事件为强终止信号"),
              ], [1500, 2760, 2760, 2340], font_size=8.7)

    add_heading(doc, "10. 分阶段上线计划", 1)
    add_table(doc,
              ["阶段", "范围", "关键验收", "是否扩量"],
              [
                  ("P0 链路验收", "测试账号与内部流量", "绑定、关注查询、Webhook、渠道白名单、+5/+15/+30 分钟任务、泰国时区窗口、短信链接、预约/出席状态和终止开关端到端通过", "全部关键链路通过才进入 P1"),
                  ("P1 关注率灰度", "少量全渠道新注册", "5 分钟查询准确；自然/电话/短信归因完整；投诉、屏蔽与技术异常在阈值内", "确认完成页和补促有效后扩量"),
                  ("P2 分配 A/B", "符合条件的非安卓用户，30%/70% 稳定随机", "组间基线平衡；销售容量可承接；付费、出席、人效和护栏指标可解释", "按预设观察窗复盘，不提前挑结果"),
                  ("P3 直分扩量", "逐步扩大非安卓直分比例", "经营指标持续优于对照且无体验/负载恶化", "30%→50%→80%→100%，每档设观察期"),
                  ("P4 OA 坐席协作", "为销售配置 OA 坐席", "机器人/真人在同一 OA 会话的锁定、接管、恢复、权限和审计通过", "达到条件后将 100% 直分设为正式机制"),
              ], [1440, 3000, 3480, 1440], font_size=8.7)

    heading = add_heading(doc, "11. 风险与决策门槛", 1)
    heading.paragraph_format.space_before = Pt(10)
    risk_table = add_table(doc,
              ["风险", "影响", "控制措施"],
              [
                  ("关注检测误判或绑定失败", "错误外呼已关注用户，损害体验并污染指标", "将技术失败与未关注分开；抽样对账；失败时不触发外呼"),
                  ("OA 与私人 LINE 重复触达", "用户困惑、屏蔽或投诉", "共享最近动作、消息频控、销售回写、结束状态强制停发"),
                  ("手机号搜索 LINE 不成功", "直分不等于成功建联", "单列 search/add 成功率；不以搜索能力作为唯一沟通路径"),
                  ("30%/70% 样本质量不平衡", "错误判断直分效果", "系统随机、固定分组、基线平衡检查、预设观察窗"),
                  ("销售容量或执行差异", "直分组积压，实验结果被执行力扭曲", "容量上限、排班、SLA、人均负载监控和培训抽检"),
                  ("自动化边界失效", "出席后仍发送课前提醒，或二次缺席后继续追踪", "状态机、结束原因、幂等测试与日常异常报表"),
                  ("通信与隐私合规", "监管、投诉与品牌风险", "泰语告知、合法同意、最小权限、保存周期、退订/屏蔽和审计"),
              ], [2340, 2580, 4440], font_size=8.1)
    for row in risk_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=55, bottom=55, start=110, end=110)
    add_callout(doc, "扩到 100% 直分的门槛", "至少同时满足：经营主指标优于抢新对照；销售人均负载可控；重复触达、OA 屏蔽与投诉没有显著恶化；数据完整且实验结果稳定；OA 坐席协作能力通过验收。", fill=PALE_GREEN, accent=GREEN)

    heading = add_heading(doc, "12. 上线验收清单", 1)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    for item in (
        "所有渠道完成页均有可用的一键关注入口，移动端与桌面端完成测试。",
        "user_id / 手机号与 LINE userId 绑定准确，可按具体用户查询关注状态。",
        "渠道白名单、泰国时区 08:00–21:00、+5/+15/+30 分钟检测与外呼、短信链接、欢迎消息均有事件、轮次和停止/跳过原因。",
        "首次预约、首次提醒、首次出席、二约、第二次提醒、最终结束形成闭环状态机。",
        "安卓、非安卓 30%、非安卓 70% 的分配规则已由系统自动执行并可审计。",
        "销售可见 OA 摘要，能回写电话打通、手机号搜索 LINE 和私人 LINE 建联结果。",
        "看板可同时查看关注漏斗、课前漏斗、分配实验、销售效率和体验护栏。",
        "终止、退订、屏蔽、技术降级、权限和异常告警已完成演练。",
    ):
        p = add_bullet(doc, item)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.95
        for run in p.runs:
            set_run_font(run, size=9.2)

    heading = add_heading(doc, "13. 官方能力依据与方案说明", 1)
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)
    add_paragraph(doc, "以下为 LINE 官方技术依据；落地仍需结合第三方接口与泰国当地合规要求评审。", size=9.2, after=2)
    sources = [
        ("LINE Login：Add friend option", "https://developers.line.biz/en/docs/line-login/link-a-bot/", "授权时展示添加 OA 选项。"),
        ("Friendship Status API", "https://developers.line.biz/en/reference/line-login/#get-friendship-status", "返回 friendFlag 判断好友状态。"),
        ("LINE URL Scheme", "https://developers.line.biz/en/docs/messaging-api/using-line-url-scheme/", "打开 OA 资料、聊天或二维码。"),
        ("Messaging API webhook", "https://developers.line.biz/en/docs/messaging-api/receiving-messages/", "接收 follow / unfollow 等事件。"),
    ]
    source_table = doc.add_table(rows=2, cols=2)
    source_table.style = "Table Grid"
    for index, (title, url, note) in enumerate(sources):
        cell = source_table.cell(index // 2, index % 2)
        set_cell_shading(cell, PALE_BLUE if index % 2 == 0 else LIGHT_GRAY)
        set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_hyperlink(p, title, url)
        r = p.add_run("｜" + note)
        set_run_font(r, size=8.5)
        set_row_cant_split(source_table.rows[index // 2])
    set_table_geometry(source_table, [4680, 4680])
    add_paragraph(doc, "注：短信深链接、触点归因、销售私人 LINE 协同和 A/B 测试设计是项目建议，不是 LINE 官方强制流程。", size=8.7, color=MID_GRAY, italic=True, after=0, before=3)

    doc.save(DOCX_PATH)


def build_markdown():
    text = f"""# 泰国 LINE-OA 课前运营与销售分配实验方案

版本：v1.0  
日期：2026-08-19  
状态：跨部门评审稿

## 核心结论

先把 LINE-OA 建成泰国所有渠道注册后的统一课前运营入口，再通过非安卓渠道 30% 直分 / 70% 抢新的随机实验验证销售归属机制。当前阶段允许 LINE-OA 与销售私人 LINE 分线沟通，但必须由 CRM 统一记录与控频；长期升级到销售 OA 坐席后，再扩大为 100% 注册即直分。

## 已确认规则

- 泰国市场，覆盖安卓与非安卓等所有注册渠道。
- 第三方工具支持 user_id / 手机号与 LINE userId 绑定，并可查询具体用户是否关注 OA。
- 注册完成页先引导关注，注册 5 分钟后查询。外呼仅对配置的较高转化渠道启用；外呼窗口暂定为泰国时间 08:00–21:00，注册和调用时点均须在窗口内，按注册后 +5 / +15 / +30 分钟最多尝试 3 次。接通或已关注即停止，超过 21:00 的尝试跳过且不顺延。
- 安卓渠道出席前不分配销售，出席后再分配。
- 非安卓 30% 注册后立即直分；OA 自动运营，销售用手机号搜索用户并以私人 LINE 单独沟通。
- 非安卓 70% 保持抢新，真人销售电话打通后才归属。
- 首次缺席后只进行 1 次二约；二约成功后继续第二次提醒；第二次出席或再次缺席后结束。

## 主流程

```mermaid
flowchart TD
  A[全渠道用户注册] --> B[完成页引导一键关注 LINE-OA]
  B --> C{{注册后 5 分钟查询关注状态}}
  C -- 已关注 --> D[绑定 user_id / 手机号 / LINE userId]
  C -- 未关注且渠道/时段符合 --> E[注册后 +5 分钟 AI 首呼]
  E -- 未接通 --> E2[+15 分钟重试]
  E2 -- 未接通 --> E3[+30 分钟重试]
  E --> F[接通后发送 OA 深链接短信]
  E2 --> F
  E3 --> F
  E3 -- 仍未接通 --> X
  F -- 后续关注成功 --> D
  F -- 仍未关注 --> X[记录原因并结束本轮 OA 自动化]
  D --> G[问好 / 轻量挖需 / FAQ]
  G --> H{{是否已预约}}
  H -- 否 --> I[OA 协助预约]
  H -- 是 --> J[首次课前提醒]
  I --> J
  J --> K{{首次是否出席}}
  K -- 是 --> L[OA SOP 结束，按渠道规则承接]
  K -- 否 --> M[发起 1 次二次预约]
  M -- 失败 --> N[记录原因并结束]
  M -- 成功 --> O[第二次课前提醒]
  O --> P{{第二次是否出席}}
  P -- 是 --> Q[OA SOP 结束，按渠道规则承接]
  P -- 否 --> R[再次缺席，OA SOP 结束]
```

## 关注率建议

1. 注册完成页移动端使用一键关注深链接，桌面端同时显示二维码。
2. 若注册/登录链路允许，优先评估 LINE Login add friend option 与 Friendship Status API。
3. 注册后 5 分钟按已绑定的用户身份查询关注状态，仅对配置渠道创建外呼任务。
4. 使用泰国时区控制外呼：窗口暂定 08:00–21:00，注册和调用时点均须在窗口内，按 +5/+15/+30 分钟最多三次；每次呼叫前重新查询，接通或已关注即停止，超窗跳过且不顺延。
5. AI 电话只说明关注价值并提示用户查收链接，不在电话中逐字符播报 OA ID。
6. 通话中或通话后发送可点击短信；OA ID 只作为深链接失败时的兜底。
7. 分别统计完成页自然关注率、各轮电话/短信补促增量、跳过原因和取消关注率。

核心口径：OA 关注率 = 统计窗口内成功绑定且 friendFlag=true 的去重注册用户数 / 符合条件的去重注册用户数。

## 销售分配 A/B 测试

```mermaid
flowchart TD
  A[泰国用户注册] --> B{{渠道类型}}
  B -- 安卓 --> C[出席前不分配，OA 负责课前运营]
  C --> D[出席后分配销售]
  B -- 非安卓 30% --> E[注册后立即直分]
  E --> F[OA 自动运营 + 销售私人 LINE 分线沟通]
  B -- 非安卓 70% --> G[进入抢新池]
  G --> H{{真人销售是否打通}}
  H -- 是 --> I[分配归属，销售可用私人 LINE]
  H -- 否 --> J[暂不归属，OA 继续课前运营]
  F --> K[长期升级 OA 销售坐席]
  I --> K
  K --> L[逐步扩大到 100% 注册即直分]
```

30%/70% 必须系统随机、用户分组固定，并按渠道、注册日期/时段检查基线平衡。销售分配实验主指标使用注册至付费转化率和每注册用户收入；OA 关注率是全项目主指标，但不是分配实验唯一判断依据。

## 跨部门职责

| 部门/角色 | 核心任务 | 意义 |
|---|---|---|
| 项目负责人/PMO | 统一范围、里程碑、风险和跨部门决策 | 保证端到端结果，而非局部工具上线 |
| 市场 | 传递渠道参数；改造完成页；稳定实验流量 | 提升自然关注并实现渠道归因 |
| 产品/产研 | 接入 OA、绑定/查询、渠道白名单、+5/+15/+30 分钟外呼、泰国时区窗口、SOP、随机分组和退出开关 | 将业务规则固化为可靠系统能力 |
| 前端运营 | 泰语话术、FAQ、预约、提醒、二约、频控和质检 | 保证自动化真正帮助用户上课 |
| 销售运营 | 归属规则、容量、私人 LINE 协同 SOP 和冲突处理 | 避免抢单争议、重复触达和负载不均 |
| 真人销售 | 执行电话/私人 LINE；查看 OA 摘要；回写结果；课后转化 | 将 AI 信息沉淀转为高价值人工跟进 |
| 数据/BI | 漏斗、实验校验、指标口径、显著性和异常监控 | 判断真实增量与因果效果 |
| 法务/合规/安全 | 通信同意、隐私、权限、保存周期、退订和审计 | 控制个人信息和通信风险 |
| 第三方供应商 | SLA、Webhook、状态查询、日志、坐席与故障恢复 | 避免关键链路成为黑盒 |

## 长期演进

1. P0：端到端链路验收。
2. P1：全渠道小流量验证自然关注与 AI 电话/短信补促。
3. P2：非安卓 30%/70% 销售分配 A/B 测试。
4. P3：按 30% → 50% → 80% → 100% 扩大直分比例。
5. P4：每位销售配置 LINE-OA 坐席，在同一 OA 会话中实现机器人/真人接管、恢复与审计。

## 官方依据

- [LINE Login add friend option](https://developers.line.biz/en/docs/line-login/link-a-bot/)
- [LINE Login Friendship Status API](https://developers.line.biz/en/reference/line-login/#get-friendship-status)
- [LINE URL Scheme](https://developers.line.biz/en/docs/messaging-api/using-line-url-scheme/)
- [Messaging API webhook](https://developers.line.biz/en/docs/messaging-api/receiving-messages/)

注：AI 电话后的短信深链接、触点归因、销售私人 LINE 协同和 A/B 测试设计属于本项目建议，并非 LINE 官方强制流程。
"""
    MD_PATH.write_text(text, encoding="utf-8")


if __name__ == "__main__":
    build_main_flow()
    build_assignment_flow()
    build_docx()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
